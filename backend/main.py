import pandas as pd
from tqdm import tqdm
import json
import os
import docx
import warnings
from dotenv import load_dotenv
from huggingface_hub import login
warnings.filterwarnings("ignore")
print("Logging into Hugging Face Hub...")
login()
print("Login successful.")

from judges.judge1_instruction_follower import get_instruction_following_score
from judges.judge2_hallucination import check_hallucination
from judges.judge3_coherence import get_coherence_score, initialize_gemini


def load_data(file_path):
    """
    Loads data from different file formats (CSV, JSON, DOCX).
    """
    _, file_extension = os.path.splitext(file_path)
    if file_extension == '.csv':
        df = pd.read_csv(file_path)
        return df.to_dict('records')
    elif file_extension == '.json':
        with open(file_path, 'r') as f:
            return json.load(f)
    elif file_extension == '.docx':
        doc = docx.Document(file_path)
        if len(doc.paragraphs) >= 2:
            prompt = doc.paragraphs[0].text
            response = "\n".join([p.text for p in doc.paragraphs[1:]])
            return [{'prompt': prompt, 'response': response}]
        else:
            print(f"Warning: DOCX file '{file_path}' has less than 2 paragraphs.")
            return []
    else:
        print(f"ERROR: Unsupported file format '{file_extension}'.")
        return None


def evaluate_agent_response(prompt: str, response: str):
    """
    The master function that runs the full evaluation pipeline.
    """
    instruction_score = get_instruction_following_score(prompt, response)
    nli_verdict, is_hallucination = check_hallucination(response)
    coherence_score = get_coherence_score(prompt, response)

    report = {
        "instruction_score_J1": round(instruction_score, 4) if instruction_score is not None else "N/A",
        "coherence_score_J3": round(coherence_score, 4),
        "is_hallucination_J2": is_hallucination,
        "nli_prediction_J2": nli_verdict
    }
    return report


if __name__ == "__main__":

    load_dotenv()
    api_key = os.getenv("GEMINI_API_KEY")
    initialize_gemini(api_key)

    input_file = "data/agent_responses_to_evaluate.csv"
    output_file = "final_evaluation_report.json"

    print(f"Loading data from {input_file}...")
    data_to_evaluate = load_data(input_file)

    if data_to_evaluate is not None:
        results = []
        print(f"Starting evaluation for {len(data_to_evaluate)} items...")

        for item in tqdm(data_to_evaluate):
            if 'prompt' in item and 'response' in item:
                report = evaluate_agent_response(item['prompt'], item['response'])
                full_result = {**item, **report}
                results.append(full_result)
            else:
                print(f" Skipping item due to missing 'prompt' or 'response': {item}")

        print(f"\n Evaluation complete. Full report saved to {output_file}")
        with open(output_file, 'w') as f:
            json.dump(results, f, indent=2)

        print("\n--- Final Report Preview (first record) ---")
        if results:
            print(json.dumps(results[0], indent=2))